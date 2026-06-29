# doc_generator.py
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_border(cell):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcPr.append(border)


def set_cell_bg(cell, color_hex="BDD7EE"):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def generate_doc(guest_data: list, output_dir: str = "output") -> str:
    """
    출연료 집행 의뢰서 Word 문서를 생성합니다.
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # ── 페이지 여백 설정
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # ── 제목
    now = datetime.now()
    title_text = f"{now.month}월 제작비(출연료) 집행 의뢰서"
    title = doc.add_paragraph(title_text)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = "맑은 고딕"

    doc.add_paragraph("")

    # ── 표 생성 (헤더 1행 + 데이터 최소 10행)
    headers = ["방송일자 (일)", "성명", "구매처/주민번호", "지급액(만원)", "연락처(H.P)/비고"]
    data_rows = max(len(guest_data), 10)
    total_rows = 1 + data_rows

    table = doc.add_table(rows=total_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 열 너비
    col_widths = [Cm(3.2), Cm(3.0), Cm(5.0), Cm(2.8), Cm(4.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    # ── 헤더 행
    header_row = table.rows[0]
    for cell, header in zip(header_row.cells, headers):
        set_cell_border(cell)
        set_cell_bg(cell, "BDD7EE")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "맑은 고딕"

    # ── 데이터 행 작성
    for i in range(data_rows):
        row = table.rows[i + 1]
        if i < len(guest_data):
            entry = guest_data[i]
            guest = entry.get("guest", "")
            date = entry.get("date", "")

            # 날짜는 있는데 이름을 못 찾은 경우 "인식실패" 표시
            is_failed = False
            if not guest and date:
                guest = "인식실패"
                is_failed = True

            row_data = [date, guest, "", "", ""]
        else:
            row_data = ["", "", "", "", ""]
            is_failed = False

        for j, (cell, value) in enumerate(zip(row.cells, row_data)):
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.font.size = Pt(10)
            run.font.name = "맑은 고딕"

            # 성명 칸이 "인식실패"이면 빨간색 굵게
            if j == 1 and is_failed:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.font.bold = True

    doc.add_paragraph("")

    # ── 서명란
    sign_lines = [
        "디지털전략국 디지털전략부 팀장 :  박 수 현  서명",
        "디지털전략국 디지털전략부 PD :  최 건 일  서명",
    ]
    for line in sign_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.name = "맑은 고딕"

    # ── 저장
    filename = f"출연료집행의뢰서_{now.strftime('%Y%m')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"✅ 문서 저장 완료: {filepath}")
    return filepath
